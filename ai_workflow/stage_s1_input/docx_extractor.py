"""Extract text and images from DOCX files.

Supports both .docx (Office Open XML) and legacy .doc (OLE) formats.
Images are extracted with a temporary generic name and returned as raw bytes
alongside the raw text content.
"""

from __future__ import annotations

import io
import os
import shutil
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import struct

try:
    from docx import Document
    from docx.oxml.ns import qn
except ImportError:
    Document = None


@dataclass
class ExtractedImage:
    """Represents an image extracted from a DOCX file."""

    index: int
    bytes_data: bytes
    content_type: str
    original_filename: Optional[str] = None
    width: Optional[int] = None
    height: Optional[int] = None


@dataclass
class ExtractedParagraph:
    """Represents a paragraph with its text and position metadata."""

    text: str
    paragraph_index: int
    style_name: Optional[str] = None
    images: list[ExtractedImage] = field(default_factory=list)


@dataclass
class ExtractedDocument:
    """Container for all extracted content from a DOCX file."""

    paragraphs: list[ExtractedParagraph] = field(default_factory=list)
    images: list[ExtractedImage] = field(default_factory=list)
    full_text: str = ""
    page_count: int = 1
    source_path: Optional[str] = None

    def get_image_count(self) -> int:
        return len(self.images)


class DocxExtractor:
    """Extract text and images from .docx files.

    For legacy .doc files, attempts conversion via LibreOffice if available.
    """

    def __init__(self):
        self._temp_dir: Optional[tempfile.TemporaryDirectory] = None

    def extract(self, docx_path: str | Path) -> ExtractedDocument:
        """Extract all text and images from a DOCX file.

        Args:
            docx_path: Path to the .docx file.

        Returns:
            ExtractedDocument containing paragraphs, images, and full text.
        """
        path = Path(docx_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {docx_path}")

        if Document is None:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )

        suffix = path.suffix.lower()
        if suffix == ".doc":
            return self._extract_from_doc_legacy(path)
        elif suffix == ".docx":
            return self._extract_from_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _extract_from_docx(self, path: Path) -> ExtractedDocument:
        doc = Document(str(path))
        result = ExtractedDocument(source_path=str(path))

        image_counter = [0]

        for para_idx, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            style_name = None
            if para.style and para.style.name:
                style_name = para.style.name

            ep = ExtractedParagraph(
                text=para_text,
                paragraph_index=para_idx,
                style_name=style_name,
            )

            for run in para.runs:
                for shape in run._element.findall(
                    ".//" "a:blip", {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
                ):
                    image_counter[0] += 1
                    img = self._extract_image_from_blip(shape, image_counter[0])
                    if img:
                        ep.images.append(img)
                        result.images.append(img)

            result.full_text += para_text + "\n"
            result.paragraphs.append(ep)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        cell_text = para.text.strip()
                        if cell_text:
                            result.full_text += cell_text + "\n"

        return result

    def _extract_image_from_blip(
        self, blip_elem, index: int
    ) -> Optional[ExtractedImage]:
        embed_attr = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        link_attr = "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}link"

        r_id = blip_elem.get(embed_attr) or blip_elem.get(link_attr)
        if not r_id:
            return None

        try:
            doc_part = blip_elem._element.nsmap.get(
                "r", "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
            )
            image_part = blip_elem._element.getparent().getparent().get(
                "TargetMode", "Internal"
            )

            rels = blip_elem._element.iter("{http://schemas.openxmlformats.org/package/2006/relationships}relationship")

            blip_parent = blip_elem._element.getparent()
            while blip_parent is not None:
                if hasattr(blip_parent, "part"):
                    doc_rels = blip_parent.part.rels
                    if r_id in doc_rels:
                        rel = doc_rels[r_id]
                        img_bytes = rel.target_part.blob
                        content_type = rel.target_part.content_type
                        ext = self._content_type_to_ext(content_type)
                        return ExtractedImage(
                            index=index,
                            bytes_data=img_bytes,
                            content_type=content_type,
                            original_filename=f"image_{index}{ext}",
                        )
                blip_parent = blip_parent.getparent()

        except Exception:
            pass

        return None

    def _content_type_to_ext(self, content_type: str) -> str:
        mapping = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/webp": ".webp",
        }
        return mapping.get(content_type, ".png")

    def _extract_from_doc_legacy(self, path: Path) -> ExtractedDocument:
        libre_office = shutil.which("libreoffice") or shutil.which("soffice")
        if not libre_office:
            raise RuntimeError(
                "Legacy .doc files require LibreOffice for conversion.\n"
                "Install LibreOffice or convert the file to .docx manually.\n"
                "On macOS: brew install --cask libreoffice"
            )

        if self._temp_dir is None:
            self._temp_dir = tempfile.TemporaryDirectory()

        temp_dir = Path(self._temp_dir.name)
        converted = temp_dir / f"{path.stem}_converted.docx"

        result = shutil.run(
            [
                libre_office,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(temp_dir),
                str(path.absolute()),
            ],
            capture_output=True,
            text=True,
            timeout=60,
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed:\n{result.stderr}"
            )

        converted_path = temp_dir / f"{path.stem}.docx"
        if not converted_path.exists():
            raise RuntimeError(
                f"Conversion output not found. Expected: {converted_path}"
            )

        return self._extract_from_docx(converted_path)

    def save_images(
        self,
        doc: ExtractedDocument,
        output_dir: str | Path,
        prefix: str = "image",
    ) -> dict[int, str]:
        """Save extracted images to a directory and return index-to-path mapping.

        Args:
            doc: The extracted document.
            output_dir: Directory to save images to.
            prefix: Filename prefix for saved images.

        Returns:
            Mapping of image index to saved file path.
        """
        out = Path(output_dir)
        out.mkdir(parents=True, exist_ok=True)

        index_to_path: dict[int, str] = {}
        for img in doc.images:
            ext = self._content_type_to_ext(img.content_type)
            filename = f"{prefix}_{img.index:03d}{ext}"
            filepath = out / filename
            filepath.write_bytes(img.bytes_data)
            index_to_path[img.index] = str(filepath)
            img.original_filename = filename

        return index_to_path

    def close(self):
        """Clean up temporary resources."""
        if self._temp_dir:
            self._temp_dir.cleanup()
            self._temp_dir = None

    def __enter__(self):
        return self

    def __exit__(self, *args):
        self.close()


def extract_docx(
    docx_path: str | Path,
    images_dir: str | Path | None = None,
) -> ExtractedDocument:
    """Convenience function for extracting text and images from a DOCX.

    Args:
        docx_path: Path to the .docx file.
        images_dir: Optional directory to save extracted images.

    Returns:
        ExtractedDocument with all content.
    """
    extractor = DocxExtractor()
    try:
        doc = extractor.extract(docx_path)
        if images_dir:
            extractor.save_images(doc, images_dir)
        return doc
    finally:
        extractor.close()
