#!/usr/bin/env python3
"""
将 Markdown 需求文档转换为 DOCX，并插入图片附件。
输出到 workflow_assets/<req_name>/「S1 需求评审」/v1.0/raw/ 目录。
"""

import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# === 配置 ===
WORKSPACE = Path("/Users/gleon/Documents/TestDev/AIDocxWorkFlow")
MD_PATH = WORKSPACE / "sample_requirements.md"
IMAGE_PATH = Path(
    "/Users/gleon/.cursor/projects/Users-gleon-Documents-TestDev-AIDocxWorkFlow/assets/"
    "0846293a-0abf-4760-93c0-84b2bf3ce571-b0648b18-57e9-4834-9c35-f0b121aa3a86.png"
)
REQ_NAME = "游戏道具商城系统"
VERSION = "v1.0"
OUTPUT_DIR = WORKSPACE / "workflow_assets" / REQ_NAME / f"「S1 需求评审」" / VERSION / "raw"
OUTPUT_DOCX = OUTPUT_DIR / f"{REQ_NAME}_{VERSION}.docx"


def set_heading_style(paragraph, level):
    """设置标题样式"""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    if level == 1:
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2F, 0x54, 0x7D)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_horizontal_rule(doc):
    """在文档中添加水平线"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)


def convert_md_to_docx(md_path: Path, docx_path: Path, image_path: Path | None = None):
    """将 Markdown 文件转换为 DOCX，可选插入图片。"""
    doc = Document()

    # 文档默认样式
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 跳过空行（由循环统一处理）
        if not stripped:
            i += 1
            continue

        # 标题行
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            p = doc.add_paragraph()
            p.add_run(title)
            set_heading_style(p, level)
            i += 1
            continue

        # 水平线 ---
        if re.match(r"^---+$", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # 列表项（无序）
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            item_text = bullet_match.group(1).strip()
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item_text)
            i += 1
            continue

        # 列表项（有序）
        numbered_match = re.match(r"^(\d+)\.\s+(.+)$", stripped)
        if numbered_match:
            item_text = numbered_match.group(2).strip()
            p = doc.add_paragraph(style="List Number")
            p.add_run(item_text)
            i += 1
            continue

        # 普通段落（可能是 markdown 链接/粗体/斜体混合）
        # 收集连续的非特殊行组成段落
        para_lines = [stripped]
        j = i + 1
        while j < len(lines) and lines[j].strip() and not re.match(r"^(#{1,3} |[-*] |\d+\. |---+$)", lines[j].strip()):
            para_lines.append(lines[j].strip())
            j += 1
        full_text = " ".join(para_lines)

        # 简单处理：移除 markdown 链接文字
        full_text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", full_text)
        full_text = re.sub(r"[*_]{1,3}([^*_]+)[*_]{1,3}", r"\1", full_text)

        p = doc.add_paragraph(full_text)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)

        i = j

    # 在最后插入图片（如果提供）
    if image_path and image_path.exists():
        doc.add_page_break()
        img_heading = doc.add_paragraph()
        img_heading.add_run("附件：参考图片").bold = True
        img_heading.runs[0].font.size = Pt(12)
        img_heading.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

        doc.add_picture(str(image_path), width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 给图片添加说明
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"（图片来源：{image_path.name}）")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.save(docx_path)
    print(f"DOCX 已保存到：{docx_path}")
    return docx_path


if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    convert_md_to_docx(MD_PATH, OUTPUT_DOCX, IMAGE_PATH)
    print(f"✅ 完成！文件：{OUTPUT_DOCX}")
